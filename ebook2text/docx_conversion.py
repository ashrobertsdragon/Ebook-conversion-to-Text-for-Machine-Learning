import base64
import logging
from typing import List, Tuple

import docx

from ._namespaces import docx_ns_map
from ._types import Document, Paragraph
from .abstract_book import (
    BookConversion,
    ChapterSplit,
    ImageExtraction,
    TextExtraction,
)
from .chapter_check import is_chapter, is_not_chapter
from .ocr import run_ocr


class DocxConverter(BookConversion):
    """
    Class to convert a Word document to structured text.

    Attributes:
        file_path (str): The path to the Word document.
        metadata (dict): Metadata related to the document.
        doc (Document): The parsed Word document object.
        paragraphs (list): List of paragraphs objects extracted from the
            document.
    """

    def __init__(self, file_path: str, metadata: dict):
        """
        Initializes the DocxConverter with file path and metadata.
        """
        super().__init__(file_path, metadata)
        self.paragraphs: list = self.extract_paragraphs()

    def _read_file(self, file_path: str) -> Document:
        """
        Reads a Word document from the specified file path.

        Returns:
            Document: The parsed Word document object.
        """
        return docx.Document(file_path)

    def extract_paragraphs(self) -> list:
        """
        Extracts paragraphs from the Word document.

        Returns:
            list: A list of paragraph objects extracted from the document.
        """
        return self.book.paragraphs

    def extract_images(self, paragraph: Paragraph) -> List[str]:
        image_extractor = DocxImageExtractor(paragraph)
        return image_extractor.extract_images()

    def extract_text(self, paragraph: Paragraph) -> str:
        text_extractor = DocxTextExtractor(paragraph, self)
        return text_extractor.extract_text()

    def split_chapters(self) -> None:
        """
        Splits the paragraphs into chapters using the ChapterSplitter.
        """
        splitter = DocxChapterSplitter(self.paragraphs, self.metadata, self)
        return splitter.split_chapters()


class DocxImageExtractor(ImageExtraction):
    """
    A class dedicated to extracting images from docx Paragraph objects.
    """

    def __init__(self, paragraph: Paragraph):
        self.paragraph = paragraph

    def extract_images(self) -> List[str]:
        """
        Extracts and converts images found in the paragraph into
        base64-encoded strings.

        Returns:
            list: A list of base64-encoded strings, each representing an image
                extracted from the paragraph.
        """
        image_blobs: list = self._extract_image_blobs()
        return self._build_base64_images_list(image_blobs)

    def _build_base64_images_list(self, image_streams: list) -> list:
        """
        Converts image blobs to base64-encoded strings.

        Args:
            image_blobs (list): A list of extracted image streams.

        Returns:
            list: A list of base64-encoded strings representing the images.
        """
        return [
            base64.b64encode(image).decode("utf-8") for image in image_streams
        ]

    def _extract_image_blobs(self) -> list:
        """
        Extracts image blobs from the paragraph.

        Returns:
            list: A list of image blobs extracted from the paragraph.
        """
        namespace: dict = docx_ns_map
        image_blobs: list = []

        blips = self.paragraph._p.findall(".//a:blip", namespaces=namespace)
        for blip in blips:
            try:
                rId = blip.attrib[f"{{{namespace['r']}}}embed"]
                image_part = self.paragraph.part.related_parts[rId]
                image_blobs.append(image_part.blob)
            except Exception as e:
                logging.exception("Could not extract images %s", str(e))
        return image_blobs


class DocxTextExtractor(TextExtraction):
    """
    Class dedicated to extracting and processing text from docx Paragraphs.
    """

    def __init__(self, paragraph: Paragraph, parent: DocxConverter):
        self.paragraph = paragraph
        self.parent = parent

    def extract_text(self) -> str:
        """
        Extracts the text content from the paragraph, performs OCR on any
        images present, and returns the combined text.
        """
        ocr_text: str = self._extract_image_text()
        paragraph_text: str = self.paragraph.text.strip()
        return ocr_text if ocr_text else paragraph_text

    def _extract_image_text(self) -> str:
        """
        Extracts text from images within the paragraph using OCR.
        """
        base64_images: list = self.parent.extract_images(self.paragraph)
        if base64_images:
            return run_ocr(base64_images)
        return ""


class DocxChapterSplitter(ChapterSplit):
    """
    Class responsible for splitting a list of paragraphs into chapters for a
    DOCX file.
    """

    def __init__(
        self,
        paragraphs: List[Paragraph],
        metadata: dict,
        parent: DocxConverter,
    ):
        super().__init__(paragraphs, metadata, parent)
        self.paragraphs = self.text_obj

        self.non_chapter: bool = False
        self.pages_list: list = []

    def split_chapters(self) -> str:
        """
        Process paragraphs to organize them into pages and chapters, handling
        page breaks and chapter starts, and compile the final structured text
        of the book.

        Returns:
            str: The structured text of the entire book.
        """
        current_page: list = []
        current_para_index: int = 0

        for paragraph in self.paragraphs:
            paragraph_text = self.parent.extract_text(paragraph)
            current_para_index += 1

            if self._contains_page_break(paragraph):
                self._add_page(current_page)
                current_page = []
                current_para_index = 0

            if paragraph_text:
                (processed_text, current_para_index) = self._process_text(
                    paragraph_text, current_para_index
                )
                current_page.append(processed_text)

        if current_page:
            self._add_page(current_page)
        return "\n".join(self.pages_list)

    def _contains_page_break(self, paragraph: Paragraph) -> bool:
        """
        Checks if a given paragraph contains a page break.
        Args:
            paragraph: The Paragraph object containing the text and formatting
        Returns:
            bool: True if the paragraph contains a page break, False otherwise
        """
        p_element = paragraph._element
        if p_element.pPr is not None:
            if p_element.pPr.pageBreakBefore is not None:
                # Explicit check required by python-dox library to avoid
                # FutureWarning warning
                return True
        return False

    def _check_index(self, index: int) -> bool:
        """
        Checks if the given index exceeds a predefined maximum limit for lines
        to check.

        Args:
            index (int): The index to be checked against the maximum limit.

        Returns:
            bool: True if the index is greater than or equal to the maximum
                limit, False otherwise.
        """
        return index >= self.MAX_LINES_TO_CHECK

    def _is_start_of_chapter(self, text: str, index: int) -> bool:
        """
        Checks if the given text marks the start of a new chapter based on
        specific criteria.

        Args:
            text (str): The text to be checked for chapter header indicators.
            index (int): The number of paragraphs since the last page break.

        Returns:
            bool: True if the text is considered the start of a chapter, False
                otherwise.
        """
        if self._check_index(index):
            return False
        return is_chapter(text)

    def _is_non_chapter(self, text: str, index: int) -> bool:
        """
        Checks if the given text is not a chapter based on specific criteria.

        Args:
            text (str): The text to be checked for being front or back matter.
            index (int): The paragraph index of the text within the chapter.

        Returns:
            bool: True if the text is not a chapter, False otherwise.
        """
        if self._check_index(index):
            return False
        return is_not_chapter(text, self.metadata)

    def _process_text(
        self, paragraph_text: str, current_para_index: int
    ) -> Tuple[str, int]:
        """
        Process a paragraph's text to determine if it starts a new chapter and
        format it accordingly.

        Args:
            paragraph_text (str): The text of the paragraph to be processed.
            current_para_index (int): The index of the paragraph
                within the current chapter.

        Returns:
            Tuple[str, int]: A tuple containing the processed text and the
                updated index of the paragraph within its chapter.
        """

        if self._is_start_of_chapter(paragraph_text, current_para_index):
            current_para_index = 0
            processed_text = self.CHAPTER_SEPARATOR if self.pages_list else ""
            self.non_chapter = False
        elif self._is_non_chapter(paragraph_text, current_para_index):
            processed_text = ""
            self.non_chapter = True
        elif self.non_chapter:
            processed_text = ""
        else:
            processed_text = self.parent.clean_text(paragraph_text)
        return processed_text, current_para_index

    def _add_page(self, current_page: list) -> None:
        """
        Adds the current page content to the `self.pages` list after filtering
        out empty pages.

        Args:
            current_page (list): The list of content for the current page.

        Effects:
            Modifies the `self.pages` list by appending the non-empty page
            content.
        """
        filtered_page = list(filter(None, current_page))
        if filtered_page:
            self.pages_list.extend(filtered_page)


def read_docx(file_path: str, metadata: dict) -> str:
    """
    Reads the contents of a DOCX file and returns the processed text.

    Args:
        file_path (str): The path to the DOCX file.
        metadata (dict): Metadata about the document.

    Returns:
        str: The processed text of the DOCX file formatted into chapters.
    """
    docx_converter = DocxConverter(file_path, metadata)
    return docx_converter.split_chapters()
