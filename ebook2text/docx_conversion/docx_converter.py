from pathlib import Path
from typing import Generator, Tuple

import docx
from docx.oxml.exceptions import XmlchemyError

from ebook2text._exceptions import DocxConversionError
from ebook2text._logger import logger
from ebook2text._types import Paragraph
from ebook2text.chapter_check import is_chapter, is_not_chapter
from ebook2text.docx_conversion.docx_text_extractor import DocxTextExtractor
from ebook2text.text_utilities import desmarten_text


class DocxConverter:
    """
    Class to convert a Word document to structured text.

    Attributes:
        file_path (str): The path to the Word document.
        metadata (dict): Metadata related to the document.
        doc (Document): The parsed Word document object.
        paragraphs (list): List of paragraphs objects extracted from the
            document.
    """

    def __init__(
        self,
        file_path: Path,
        metadata: dict,
        text_extractor: DocxTextExtractor,
    ):
        """
        Initializes the DocxConverter with file path and metadata.
        """
        self.paragraphs: Generator[Paragraph, None, None] = self._read_file(
            file_path
        )
        self.metadata = metadata
        self.text_extractor = text_extractor

        self._chapter_separator: str = "***"
        self._max_lines_to_check: int = 6

        self.non_chapter: bool = False

    def _read_file(self, file_path: Path) -> Generator[Paragraph, None, None]:
        """
        Reads a Word document from the specified file path.

        Returns:
            Generator[Paragraph]: A generator of paragraph objects extracted
                from the Word document object.
        """
        try:
            yield from docx.Document(file_path).paragraphs
        except (OSError, ValueError, XmlchemyError) as e:
            logger.error(f"Error reading Word document: {e}")
            raise DocxConversionError from e

    def parse_file(self) -> Generator[str, None, None]:
        """
        Process paragraphs to organize them into pages and chapters, handling
        page breaks and chapter starts, and compile the final structured text
        of the book.

        Returns:
            str: The structured text of the entire book.
        """
        current_page: list = []
        current_para_index: int = 0

        for paragraph in self.paragraphs:
            paragraph_text = self.text_extractor.extract_text(paragraph)
            current_para_index += 1

            if self._contains_page_break(paragraph):
                if current_page:
                    yield "\n".join(current_page)
                current_page = []
                current_para_index = 0

            if paragraph_text:
                (processed_text, current_para_index) = self._process_text(
                    paragraph_text, current_para_index
                )
                if processed_text:
                    current_page.append(processed_text)

        if current_page:
            yield "\n".join(current_page)

    def _clean_before_write(self, text: str, output_path: Path) -> str:
        """
        Strips the chapter separator from the text if the file does not exist.

        Args:
            text (str): The text to be cleaned.
            output_path (Path): The path to the output file.

        Returns:
            str: The cleaned text.

        Note:
            This method is used to strip the leading chapter separator
            before writing to a file.
        """
        return (
            text
            if output_path.exists()
            else text.lstrip(self._chapter_separator)
        )

    def write_text(self, content: str, output_path: Path) -> None:
        """
        Write the parsed text to a file.

        Args:
            output (str): The parsed text to be written to the file.
            file_path (Path): The path to the output file.
        """
        cleaned_content = self._clean_before_write(content, output_path)
        with output_path.open("a", encoding="utf-8") as f:
            f.write(cleaned_content + "\n")

    def return_string(self, generator: Generator[str, None, None]) -> str:
        """
        Return the parsed text as a string.

        Args:
            generator (Generator[str, None, None]): The content generator
                that yields the text. This is usually the `parse_file` method.

        Returns:
            str: The parsed text as a single string.
        """
        return "\n".join(line for line in generator if line.strip()).lstrip(
            self._chapter_separator
        )

    @staticmethod
    def _remove_smart_punctuation(text: str) -> str:
        """
        Remove smart punctuation from the given text.

        Args:
            text (str): The text to have smart punctuation removed.

        Returns:
            str: The text with smart punctuation removed.
        """
        return desmarten_text(text)

    def _contains_page_break(self, paragraph: Paragraph) -> bool:
        """
        Checks if a given paragraph contains a page break.
        Args:
            paragraph: The Paragraph object containing the text and formatting
        Returns:
            bool: True if the paragraph contains a page break, False otherwise
        """
        p_element = paragraph._element
        return (
            p_element.pPr is not None
            and p_element.pPr.pageBreakBefore is not None
        )

    def _check_index(self, index: int) -> bool:
        """
        Checks if the given index exceeds a predefined maximum limit for lines
        to check.

        Args:
            index (int): The index to be checked against the maximum limit.

        Returns:
            bool: True if the index is greater than or equal to the maximum
                limit, False otherwise.
        """
        return index >= self._max_lines_to_check

    def _is_start_of_chapter(self, text: str, index: int) -> bool:
        """
        Checks if the given text marks the start of a new chapter based on
        specific criteria.

        Args:
            text (str): The text to be checked for chapter header indicators.
            index (int): The number of paragraphs since the last page break.

        Returns:
            bool: True if the text is considered the start of a chapter, False
                otherwise.
        """
        return False if self._check_index(index) else is_chapter(text)

    def _is_non_chapter(self, text: str, index: int) -> bool:
        """
        Checks if the given text is not a chapter based on specific criteria.

        Args:
            text (str): The text to be checked for being front or back matter.
            index (int): The paragraph index of the text within the chapter.

        Returns:
            bool: True if the text is not a chapter, False otherwise.
        """
        if self._check_index(index):
            return False
        return is_not_chapter(text, self.metadata)

    def _process_text(
        self, paragraph_text: str, current_para_index: int
    ) -> Tuple[str, int]:
        """
        Process a paragraph's text to determine if it starts a new chapter and
        format it accordingly.

        Args:
            paragraph_text (str): The text of the paragraph to be processed.
            current_para_index (int): The index of the paragraph
                within the current chapter.

        Returns:
            Tuple[str, int]: A tuple containing the processed text and the
                updated index of the paragraph within its chapter.
        """

        if self._is_start_of_chapter(paragraph_text, current_para_index):
            current_para_index = 0
            processed_text = self._chapter_separator
            self.non_chapter = False
        elif self._is_non_chapter(paragraph_text, current_para_index):
            processed_text = ""
            self.non_chapter = True
        elif self.non_chapter:
            processed_text = ""
        else:
            processed_text = self._remove_smart_punctuation(paragraph_text)
        return processed_text, current_para_index
