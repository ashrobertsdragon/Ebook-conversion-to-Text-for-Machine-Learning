import pytest
from docx import Document  # constructor
from docx.oxml import parse_xml
from docx.text.paragraph import Paragraph

from ebook2text.docx_conversion import (
    DocxConverter,
    DocxImageExtractor,
    DocxTextExtractor,
)
from ebook2text.docx_conversion._namespaces import docx_ns_map


@pytest.fixture(scope="session")
def docx_file(test_files_dir):
    """Fixture to provide path to base test docx file for testing."""
    doc_path = test_files_dir / "test_docx.docx"
    return str(doc_path)


@pytest.fixture(scope="session")
def docx_file_with_image(test_files_dir):
    """Fixture to provide path to test docx file with image for testing."""
    doc_path = test_files_dir / "test_docx_with_image.docx"
    return str(doc_path)


@pytest.fixture
def docx_paragraph_with_image(docx_file_with_image):
    """Fixture to provide a paragraph with an image for testing."""
    doc = Document(docx_file_with_image)
    return doc.paragraphs[8]


@pytest.fixture
def docx_paragraph_without_image(docx_file):
    """Fixture to provide a paragraph without an image for testing."""
    doc = Document(docx_file)
    return doc.paragraphs[0]


@pytest.fixture
def docx_image_extractor():
    """Fixture to provide an instance of DocxImageExtractor for testing."""
    return DocxImageExtractor()


@pytest.fixture
def docx_image_extractor_with_image(
    docx_image_extractor, docx_paragraph_with_image
):
    """Fixture to provide an instance of DocxImageExtractor for testing."""
    extractor = docx_image_extractor
    extractor.paragraph = docx_paragraph_with_image
    return extractor


@pytest.fixture
def docx_text_extractor(docx_image_extractor):
    return DocxTextExtractor(docx_image_extractor)


@pytest.fixture
def docx_text_extractor_with_image(
    docx_text_extractor, docx_paragraph_with_image
):
    return DocxTextExtractor(docx_paragraph_with_image)


@pytest.fixture
def docx_converter(docx_file, metadata, docx_text_extractor):
    """Fixture to initialize DocxConverter with a test DOCX file and metadata."""
    return DocxConverter(docx_file, metadata, docx_text_extractor)


@pytest.fixture
def docx_converter_with_image(
    docx_file_with_image, metadata, docx_text_extractor_with_image
):
    """Fixture to initialize DocxConverter with a test DOCX file with image and metadata."""
    return DocxConverter(
        docx_file_with_image, metadata, docx_text_extractor_with_image
    )


def test_read_file(docx_converter):
    """Test that a DOCX file is read and parsed correctly."""
    paragraphs = list(docx_converter.paragraphs)
    assert isinstance(paragraphs[0], Paragraph)
    assert paragraphs


def test_extract_images_with_image(
    docx_image_extractor, docx_paragraph_with_image, expected_base64_image
):
    """
    Test that `extract_images` returns a list with a base64-encoded string
    when the paragraph contains an image.
    """

    base64_images = docx_image_extractor.extract_images(
        docx_paragraph_with_image
    )

    assert len(base64_images) == 1
    assert base64_images[0] == expected_base64_image


def test_extract_images_no_image(
    docx_image_extractor, docx_paragraph_without_image
):
    """
    Test that `extract_images` returns an empty list when the paragraph
    contains no images.
    """
    base64_images = docx_image_extractor.extract_images(
        docx_paragraph_without_image
    )
    assert len(base64_images) == 0


def test_build_base64_images_list(
    docx_image_extractor_with_image,
    docx_paragraph_with_image,
    expected_base64_image,
):
    """
    Test `_build_base64_images_list` method directly to ensure it converts
    image blobs to base64-encoded strings correctly.
    """
    image_blobs = docx_image_extractor_with_image._extract_image_blobs(
        docx_paragraph_with_image
    )  # Get raw image data
    base64_images = docx_image_extractor_with_image._build_base64_images_list(
        image_blobs
    )

    assert len(image_blobs) == 1
    assert base64_images[0] == expected_base64_image


def test_extract_images_from_valid_paragraph(
    docx_image_extractor_with_image, docx_paragraph_with_image
):
    """Test successful extraction from a paragraph known to contain an image"""
    blobs = docx_image_extractor_with_image._extract_image_blobs(
        docx_paragraph_with_image
    )

    assert len(blobs) > 0
    assert isinstance(blobs[0], bytes)


def test_extract_images_from_paragraph_without_images(
    docx_image_extractor, docx_paragraph_without_image
):
    """Test extraction from a paragraph known to not contain images"""
    blobs = docx_image_extractor._extract_image_blobs(
        docx_paragraph_without_image
    )

    assert blobs == []


def test_missing_embed_attribute(
    docx_image_extractor, docx_paragraph_with_image, mocker
):
    """Test handling of missing embed attribute in blip"""
    mock_findall = mocker.patch.object(docx_paragraph_with_image._p, "findall")
    # Create a malformed blip without embed attribute
    blip = parse_xml("""
        <a:blip xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
        </a:blip>
    """)
    mock_findall.return_value = [blip]
    blobs = docx_image_extractor._extract_image_blobs(
        docx_paragraph_with_image
    )
    assert blobs == []


def test_invalid_relationship_id(docx_image_extractor, mocker):
    """Test handling of invalid relationship ID"""
    mock_paragraph = mocker.Mock()

    mock_blip = mocker.Mock()
    mock_blip.attrib = {f"{{{docx_ns_map['r']}}}embed": "rId123"}
    mock_paragraph._p = mocker.Mock()
    mock_paragraph._p.findall.return_value = [mock_blip]

    mock_part = mocker.Mock()
    mock_part.related_parts = {}
    mock_paragraph.part = mock_part
    # Leave related_parts empty so any rId will be invalid
    blobs = docx_image_extractor._extract_image_blobs(mock_paragraph)
    assert blobs == []


def test_malformed_paragraph_structure(docx_image_extractor, mocker):
    """Test handling of malformed paragraph structure"""
    mock_paragraph = mocker.Mock()
    mock_paragraph._p = mocker.Mock()
    mock_paragraph._p.findall.return_value = []
    mock_paragraph.part = None
    blobs = docx_image_extractor._extract_image_blobs(mock_paragraph)
    assert blobs == []


def test_corrupted_image_data(docx_image_extractor, mocker):
    """Test handling of corrupted image data"""
    mock_paragraph = mocker.Mock()

    mock_blip_valid = mocker.Mock()
    rId_valid = "rId123"
    mock_blip_valid.attrib = {f"{{{docx_ns_map['r']}}}embed": rId_valid}

    mock_blip_corrupt = mocker.Mock()
    rId_corrupt = "rIdCorrupt"
    mock_blip_corrupt.attrib = {f"{{{docx_ns_map['r']}}}embed": rId_corrupt}

    mock_paragraph._p = mocker.Mock()
    mock_paragraph._p.findall.return_value = [
        mock_blip_valid,
        mock_blip_corrupt,
    ]

    mock_part = mocker.Mock()
    mock_paragraph.part = mock_part
    mock_image_part_valid = mocker.Mock()
    mock_image_part_valid.blob = b"valid_image_data"

    mock_image_part_corrupt = mocker.Mock()
    type(mock_image_part_corrupt).blob = mocker.PropertyMock(
        side_effect=ValueError("Corrupted image data")
    )

    mock_part.related_parts = {
        rId_valid: mock_image_part_valid,
        rId_corrupt: mock_image_part_corrupt,
    }

    blobs = docx_image_extractor._extract_image_blobs(mock_paragraph)
    print("Extracted blobs:", blobs)
    assert len(blobs) == 1, f"Unexpected blobs: {blobs}"
    assert blobs[0] == b"valid_image_data"
