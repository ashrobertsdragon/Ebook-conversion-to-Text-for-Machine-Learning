from unittest.mock import Mock, patch

import pytest
from docx import Document
from docx.oxml import parse_xml

from ebook2text._namespaces import docx_ns_map
from ebook2text.docx_conversion import (
    DocxChapterSplitter,
    DocxConverter,
    DocxImageExtractor,
    DocxTextExtractor,
)
from ebook2text.ocr import run_ocr


@pytest.fixture(scope="session")
def docx_file(test_files_dir):
    """Fixture to provide path to base test docx file for testing."""
    doc_path = test_files_dir / "test_docx.docx"
    return str(doc_path)


@pytest.fixture(scope="session")
def docx_file_with_image(test_files_dir):
    """Fixture to provide path to test docx file with image for testing."""
    doc_path = test_files_dir / "test_docx_with_image.docx"
    return str(doc_path)


@pytest.fixture
def docx_converter(docx_file, metadata):
    """Fixture to initialize DocxConverter with a test DOCX file and metadata."""
    return DocxConverter(docx_file, metadata)


@pytest.fixture
def docx_converter_with_image(docx_file_with_image, metadata):
    """Fixture to initialize DocxConverter with a test DOCX file with image and metadata."""
    return DocxConverter(docx_file_with_image, metadata)


@pytest.fixture
def docx_paragraph_with_image(docx_file_with_image):
    """Fixture to provide a paragraph with an image for testing."""
    doc = Document(docx_file_with_image)
    return doc.paragraphs[1]


@pytest.fixture
def docx_paragraph_without_image(docx_file):
    """Fixture to provide a paragraph without an image for testing."""
    doc = Document(docx_file)
    return doc.paragraphs[0]


@pytest.fixture
def docx_image_extractor():
    """Fixture to provide an instance of DocxImageExtractor for testing."""
    return DocxImageExtractor()


@pytest.fixture
def docx_image_extractor_with_image(
    docx_image_extractor, docx_paragraph_with_image
):
    """Fixture to provide an instance of DocxImageExtractor for testing."""
    extractor = docx_image_extractor
    extractor.paragraph = docx_paragraph_with_image
    return extractor


@pytest.fixture
def expected_base64_image():
    """Fixture to provide the expected base64-encoded string of an image."""
    return "your_base64_encoded_string_here"


def test_read_file(docx_converter):
    """Test that a DOCX file is read and parsed correctly."""
    doc = docx_converter._read_file(docx_converter.file_path)
    assert isinstance(doc, Document)
    assert len(doc.paragraphs) > 0


def test_extract_images_with_image(
    docx_image_extractor, docx_paragraph_with_image, expected_base64_image
):
    """
    Test that `extract_images` returns a list with a base64-encoded string
    when the paragraph contains an image.
    """

    base64_images = docx_image_extractor.extract_images(
        docx_paragraph_with_image
    )

    assert len(base64_images) == 1
    assert base64_images[0] == expected_base64_image


def test_extract_images_no_image(
    docx_image_extractor, docx_paragraph_without_image
):
    """
    Test that `extract_images` returns an empty list when the paragraph
    contains no images.
    """
    base64_images = docx_image_extractor.extract_images(
        docx_paragraph_without_image
    )
    assert len(base64_images) == 0


def test_build_base64_images_list(
    docx_image_extractor_with_image, expected_base64_image
):
    """
    Test `_build_base64_images_list` method directly to ensure it converts
    image blobs to base64-encoded strings correctly.
    """
    image_blobs = (
        docx_image_extractor_with_image._extract_image_blobs()
    )  # Get raw image data
    base64_images = docx_image_extractor_with_image._build_base64_images_list(
        image_blobs
    )

    assert len(image_blobs) == 1
    assert base64_images[0] == expected_base64_image


def test_extract_images_from_valid_paragraph(docx_image_extractor_with_image):
    """Test successful extraction from a paragraph known to contain an image"""
    blobs = docx_image_extractor_with_image._extract_image_blobs()

    assert len(blobs) > 0
    assert isinstance(blobs[0], bytes)


def test_extract_images_from_paragraph_without_images(
    docx_image_extractor, docx_paragraph_without_image
):
    """Test extraction from a paragraph known to not contain images"""
    docx_image_extractor.paragraph = docx_paragraph_without_image
    blobs = docx_image_extractor._extract_image_blobs()

    assert blobs == []


def test_missing_embed_attribute(
    docx_image_extractor, docx_paragraph_with_image
):
    """Test handling of missing embed attribute in blip"""
    with patch.object(docx_paragraph_with_image._p, "findall") as mock_findall:
        # Create a malformed blip without embed attribute
        blip = parse_xml("""
            <a:blip xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
            </a:blip>
        """)
        mock_findall.return_value = [blip]
        docx_image_extractor.paragraph = docx_paragraph_with_image
        blobs = docx_image_extractor._extract_image_blobs()
        assert blobs == []


def test_invalid_relationship_id(
    docx_image_extractor, docx_paragraph_with_image
):
    """Test handling of invalid relationship ID"""
    with patch.object(
        docx_paragraph_with_image.part, "related_parts", new_callable=dict
    ):
        # Leave related_parts empty so any rId will be invalid
        docx_image_extractor.paragraph = docx_paragraph_with_image
        blobs = docx_image_extractor._extract_image_blobs()
        assert blobs == []


def test_malformed_paragraph_structure(
    docx_image_extractor, docx_paragraph_with_image
):
    """Test handling of malformed paragraph structure"""
    with patch.object(docx_paragraph_with_image, "part", None):
        docx_image_extractor.paragraph = docx_paragraph_with_image
        blobs = docx_image_extractor._extract_image_blobs()
        assert blobs == []


def test_corrupted_image_data(docx_paragraph_with_image):
    """Test handling of corrupted image data"""
    with patch.dict(
        docx_paragraph_with_image.part.related_parts
    ) as mock_related_parts:
        # Get the first image relationship ID
        blips = docx_paragraph_with_image._p.findall(
            ".//a:blip", namespaces=docx_ns_map
        )
        if not blips:
            pytest.skip("No images found in test document")

        rId = blips[0].attrib[f"{{{docx_ns_map['r']}}}embed"]

        # Replace the actual image part with one that raises ValueError
        mock_image_part = Mock()
        mock_image_part.blob = property(
            lambda self: exec('raise ValueError("Corrupted image data")')
        )
        mock_related_parts[rId] = mock_image_part

        extractor = DocxImageExtractor(docx_paragraph_with_image)
        blobs = extractor._extract_image_blobs()

        assert len(blobs) == len(blips) - 1


@pytest.mark.parametrize(
    "exception_type",
    [
        KeyError("Missing embed attribute"),
        KeyError("Invalid relationship ID"),
        AttributeError("Missing part attribute"),
        ValueError("Corrupted image data"),
    ],
)
def test_exception_logging(
    docx_image_extractor, docx_paragraph_with_image, exception_type, caplog
):
    """Test that exceptions are properly logged"""
    with patch.object(
        DocxImageExtractor, "_extract_image_blobs"
    ) as mock_extract:
        mock_extract.side_effect = exception_type
        docx_image_extractor.paragraph = docx_paragraph_with_image
        docx_image_extractor._extract_image_blobs()

        assert "Corrupted image data" in caplog.text


def test_extract_paragraphs(docx_converter):
    """Test paragraph extraction from the DOCX document."""
    paragraphs = docx_converter.extract_paragraphs()
    assert len(paragraphs) == 5


def test_extract_text(docx_converter):
    """Test text extraction from paragraphs."""
    paragraphs = docx_converter.extract_paragraphs()
    text_extractor = DocxTextExtractor(docx_converter)
    text = text_extractor.extract_text(paragraphs[0])
    assert text == "This is the introduction paragraph."


def test_split_chapters(docx_converter):
    """Test splitting paragraphs into chapters."""
    structured_text = docx_converter.split_chapters()
    assert "Chapter 1" in structured_text
    assert "Chapter 2" in structured_text


def test_extract_image_text_no_images(docx_converter):
    """Test image extraction when no images are present in paragraphs."""
    paragraphs = docx_converter.extract_paragraphs()
    text_extractor = DocxTextExtractor(docx_converter)
    extracted_text = text_extractor.extract_text(paragraphs[0])
    assert extracted_text == ""


def test_extract_images_with_mock_ocr(docx_converter_with_image, monkeypatch):
    """Test OCR and image extraction with a mock run_ocr function."""
    paragraphs = docx_converter.extract_paragraphs()
    monkeypatch.setattr(run_ocr, "run_ocr", lambda _: "Mocked OCR Text")

    text_extractor = DocxTextExtractor(docx_converter)
    extracted_text = text_extractor.extract_text(paragraphs[0])

    assert extracted_text == "Mocked OCR Text"


def test_chapter_splitter(docx_converter):
    """Test the chapter splitting logic for correct chapter organization."""
    paragraphs = docx_converter.extract_paragraphs()
    chapter_splitter = DocxChapterSplitter(
        paragraphs, docx_converter.metadata, docx_converter
    )
    structured_text = chapter_splitter.split_chapters()

    assert "Chapter 1" in structured_text
    assert "Chapter 2" in structured_text


def test_check_index(docx_converter):
    """Test the index check limit in chapter splitter."""
    chapter_splitter = DocxChapterSplitter(
        [], docx_converter.metadata, docx_converter
    )
    assert chapter_splitter._check_index(100) is False
    assert (
        chapter_splitter._check_index(chapter_splitter.MAX_LINES_TO_CHECK)
        is True
    )
