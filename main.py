import base64
import logging
import os
from io import BytesIO
from typing import Any, Tuple

import docx
import ebooklib
import requests
from dotenv import load_dotenv
from bs4 import BeautifulSoup
from ebooklib import epub
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import resolve1
from pdfminer.pdftypes import PDFStream
from pdfminer.layout import LTChar, LTContainer, LTText
from pdfminer.high_level import extract_pages
from PIL import Image


load_dotenv()

error_logger = logging.getLogger('error_logger')
error_logger.setLevel(logging.ERROR)
error_handler = logging.FileHandler('error.log')
error_formatter = logging.Formatter('%(asctime)s - %(name)s - ERROR - %(message)s')
error_handler.setFormatter(error_formatter)
error_logger.addHandler(error_handler)

info_logger = logging.getLogger('info_logger')
info_logger.setLevel(logging.INFO)
info_handler = logging.FileHandler('info.log')
info_formatter = logging.Formatter('%(asctime)s - %(name)s - INFO - %(message)s')
info_handler.setFormatter(info_formatter)
info_logger.addHandler(info_handler)

NOT_CHAPTER = [
    "about", "acknowledgements", "afterward", "annotation", "appendix",
    "assessment", "backmatter", "bibliography", "colophon", "conclusion",
    "contents", "contributors", "copyright", "cover", "credits", "dedication",
    "division", "endnotes", "epigraph", "errata", "footnotes", "forward",
    "frontmatter", "glossary", "imprintur", "imprint", "index", "introduction",
    "landmarks", "list", "notice", "page", "preamble", "preface", "prologue",
    "question", "rear", "revision", "sign up", "table", "toc", "volume",
    "warning"
]


def read_text_file(file: str) -> str:
  with open(file, "r") as f:
    read_file = f.read()
  return read_file


def write_to_file(content: str, file: str):
  with open(file, "w") as f:
    f.write(content, file)


def roman_to_int(roman: str) -> int:
  """
  Convert a Roman numeral to an integer.
  Arguments:
    roman (str): A string representing the Roman numeral.
  Returns int: The integer value of the Roman numeral.
  """

  roman_numerals = {
      'I': 1,
      'V': 5,
      'X': 10,
      'L': 50,
      'C': 100,
      'D': 500,
      'M': 1000
  }
  roman = roman.upper()
  total = 0
  prev_value = 0
  for char in reversed(roman):
    if char not in roman_numerals:
      raise ValueError("not Roman numeral")
    value = roman_numerals[char]
    if value < prev_value:
      total -= value
    else:
      total += value
    prev_value = value
  if not total:
    raise ValueError("Not roman numeral")
  return total


def word_to_num(number_str: str) -> int:
  """
  Convert a spelled-out number without spaces or hyphens to a string of the 
  integer. The function supports numbers from 0 to 99 and is case insensitive.
  Arguments:
    number_str (str): A string representing the spelled-out number.
  Returns int: The integer value of the spelled-out number.
  """

  num_words = {
      'zero': 0,
      'one': 1,
      'two': 2,
      'three': 3,
      'four': 4,
      'five': 5,
      'six': 6,
      'seven': 7,
      'eight': 8,
      'nine': 9,
      'ten': 10,
      'eleven': 11,
      'twelve': 12,
      'thirteen': 13,
      'fourteen': 14,
      'fifteen': 15,
      'sixteen': 16,
      'seventeen': 17,
      'eighteen': 18,
      'nineteen': 19,
      'twenty': 20,
      'thirty': 30,
      'forty': 40,
      'fifty': 50,
      'sixty': 60,
      'seventy': 70,
      'eighty': 80,
      'ninety': 90
  }
  number_str = number_str.lower()
  total = 0
  temp_word = ""
  for char in number_str:
    temp_word += char
    if temp_word in num_words:
      total += num_words[temp_word]
      temp_word = ""
  if temp_word:
    raise ValueError(f"Unknown number word: {temp_word}")
  return total


def is_chapter(s: str) -> bool:
  """
  Check if a string contains the word 'chapter', a Roman numeral, a 
  spelled-out number, or a digit.
  Arguments:
    s (str): The string to check.
  Returns bool: True if the string meets the criteria, False otherwise.
  """

  def is_number(s: str) -> bool:
    """
    Check if a string contains a number value either as digits,
    roman numerals, or spelled-out numbers.
    """

    return s.isdigit() or is_roman_numeral(s) or is_spelled_out_number(s)

  def is_roman_numeral(word: str) -> bool:
    """
    Try to convert the word to an integer. If it's not a valid roman numeral,
    it will raise a ValueError
    """

    try:
      roman_to_int(word)
      return True
    except ValueError:
      return False

  def is_spelled_out_number(s: str) -> bool:
    """
    Try to convert the word to an integer. If it's not a valid spelled-out number,
    it will raise a ValueError
    """

    try:
      word_to_num(s)
      return True
    except ValueError:
      return False

  lower_s = s.lower()
  return lower_s.startswith("chapter") or (len(lower_s.split()) == 1
                                           and is_number(lower_s))


def is_not_chapter(paragraph: str, metadata: dict) -> bool:
  """
  Checks if the given line is not a chapter.
  """
  title = metadata.get("title", "no title found")
  author = metadata.get("author", "no author found")
  paragraph = paragraph.lower()
  return any(
      paragraph.startswith(title.lower()) or paragraph.startswith(
          author.lower()) or paragraph.startswith(not_chapter_word)
      for not_chapter_word in NOT_CHAPTER)


def desmarten_text(book_content: str) -> str:
  """
  Replace smart punctuation in the given text with regular ASCII punctuation.
  Arguments:
    book_content (str): The input text containing smart punctuation.
  Returns str: The text with smart punctuation replaced by regular punctuation.
  """

  punctuation_map = str.maketrans({
      "‘": "'",
      "’": "'",
      "“": '"',
      "”": '"',
      "–": "-",
      "—": "-",
      "…": "...",
      "•": "*"
  })
  return book_content.translate(punctuation_map)


def encode_image(image_path):
  with open(image_path, "rb") as image_file:
    return base64.b64encode(image_file.read()).decode('utf-8')


def run_ocr(base64_images: list) -> str:
  """
  Perform optical character recognition (OCR) on a list of base64-encoded 
  images using the OpenAI API.

  Arguments:
    base64_images (list): A list of base64-encoded images.

  Returns str: The recognized text from the images.
  """

  api_key = os.getenv("OPENAI_API_KEY")
  image_role_list = []
  headers = {
      "Content-Type": "application/json",
      "Authorization": f"Bearer {api_key}"
  }
  for base64_image in base64_images:
    image_role_list.append({
        "type": "image_url",
        "image_url": {
            "url": f"data:image/jpeg;base64,{base64_image}",
            "detail": "low"
        }
    })
  payload = {
      "model":
      "gpt-4-vision-preview",
      "messages": [{
          "role":
          "user",
          "content": [{
              "type":
              "text",
              "text":
              ("Please provide the text in these images as a single combined "
               "statement with spaces as appropriate without any commentary. Use "
               "your judgment on whether consecutive images are a single word or "
               "multiple words. If there is no text in the image, or it is "
               "unreadable, respond with 'No text found'")
          }, *image_role_list]
      }],
      "max_tokens":
      10
  }
  try:
    response = requests.post("https://api.openai.com/v1/chat/completions",
                             headers=headers,
                             json=payload)
    if response.json()["choices"]:
      answer = response.json()["choices"][0]["message"]["content"]
      return answer
    if response.json()["error"]:
      error = response.json()["error"]["message"]
      raise Exception(error)
  except KeyError as e:
    print(f"KeyError: {e}")
    print("Response:", response.text)
    return None


def _extract_epub_chapter_text(item, book) -> str:
  """
  Extracts text from a chapter item.
  Arguments:
    item: ebooklib item representing a chapter.
  Returns string containing the text of the chapter.
  """

  soup = BeautifulSoup(item.content, "html.parser")
  elements = soup.find_all(["p", "img", "h1", "h2", "h3", "h4", "h5", "h6"])
  for i, element in enumerate(elements[:3]):
    if element.name == "img":
      image_data = book.read_item(element["src"])
      base64_images = [encode_image(image_data)]
      text = run_ocr(base64_images)
    else:
      text = element.get_text().strip().lower()
    if any(non_chapter_word in text for non_chapter_word in NOT_CHAPTER):
      return ""
    elif is_chapter(text):
      starting_line = i + 1
      return "\n".join(tag.get_text().strip()
                       for tag in elements[starting_line:] if tag != "img")
  return ""


def read_epub(file_path: str, metadata: dict) -> str:
  """
  Reads the contents of an epub file and returns it as a string.
  Arguments:
    file_path: Path to the epub file.
  Returns the contents of the epub file as a string.
  """

  chapter_contents = []
  book = epub.read_epub(file_path, options={"ignore_ncx": True})
  for item in book.get_items():
    if item.get_type() == ebooklib.ITEM_DOCUMENT and not is_not_chapter(
        item.file_name.lower(), metadata):
      chapter_text = _extract_epub_chapter_text(item, book)
      if chapter_text:
        chapter_text = desmarten_text(chapter_text)
        chapter_contents.append(chapter_text)
  return "\n***\n".join(chapter_contents)


def _docx_contains_page_break(paragraph: str) -> bool:
  """
  Checks if the given paragraph contains a page break.
  Arguments:
    paragraph: A paragraph from a docx document.

  Returns True if a page break is found, False otherwise.
  """

  for run in paragraph.runs:
    if "<w:lastRenderedPageBreak/>" in run._element.xml or "<w:br " in run.element.xml:
      return True
  return False


def _docx_extract_images(paragraph, doc) -> str:
  for run in paragraph.runs:
    for inline_shape in run.element.findall(".//a:blip",
                                            namespaces=docx.oxml.ns.nsmap):
      image_blip = inline_shape.get(
          '{http://schemas.openxmlformats.org/drawingml/2006/main}embed')
      image_part = doc.part.related_parts[image_blip]
      stream = BytesIO(image_part.blob)
      base64_images = [base64.b64encode(stream.read().decode("utf-8"))]
      ocr_text = run_ocr(base64_images)
      if ocr_text:
        return ocr_text
      else:
        return ""


def read_docx(file_path: str, metadata: dict) -> str:
  """
  Reads the contents of a docx file and returns it as a string.
  Arguments:
    file_path: Path to the docx file.
  Returns the contents of the docx file as a string.
  """

  current_page = []
  pages = []
  line_counter = 0
  max_lines_to_check = 3
  chapter_header = False
  not_chapter_flag = False
  doc = docx.Document(file_path)
  for paragraph in doc.paragraphs:
    if line_counter < max_lines_to_check:
      ocr_text = _docx_extract_images(paragraph)
    paragraph_text = ocr_text if ocr_text else paragraph.text.strip()
    if _docx_contains_page_break(paragraph) and current_page:
      pages.append("\n".join(current_page))
      chapter_header = False
      not_chapter_flag = False
      current_page = []
      line_counter = 0
    if not paragraph_text:
      continue
    elif is_chapter(paragraph_text):
      chapter_header = True
      not_chapter_flag = False
      if pages:
        paragraph_text = "***"
      else:
        continue
    elif line_counter < max_lines_to_check and not chapter_header and is_not_chapter(
        paragraph_text, metadata):
      not_chapter_flag = True
      current_page = []
      continue
    elif not_chapter_flag is True:
      continue
    paragraph_text = desmarten_text(paragraph_text)
    current_page.append(paragraph_text)
    line_counter += 1
  if current_page:
    pages.append("\n".join(current_page))
  return "\n".join(pages)



def create_image_from_binary(binary_data, width: int, height: int) -> str:
  """
  Create an image from binary data.

  Arguments:
    binary_data (bytes): The binary data of the image.
    width (int): The width of the image.
    height (int): The height of the image.

  Returns str: The string representation of the image.
  """
  try:
    image = Image.frombytes('1', (width, height), binary_data)
    image = image.convert('L')
    buffered = BytesIO()
    image.save(buffered, format="JPEG")
    image = base64.b64encode(buffered.getvalue()).decode('utf-8')
    return image
  except Exception as e:
    error_logger.exception(f"failed to create base64 encoded image due to {e}")


def extract_image(document: bytes, obj_num: int, attempt: int) -> str:
  """
  Convert binary image data into a base64-encoded JPEG string.

  This function takes binary data of an image, along with its width and 
  height, and converts it into a base64-encoded JPEG format string. This is 
  useful for encoding images in a format that can be easily transmitted or 
  stored.

  Arguments:
    binary_data (bytes): The binary data of the image.
    width (int): The width of the image in pixels.
    height (int): The height of the image in pixels.

  Returns str: A base64-encoded string representing the JPEG image.

  Raises:
    Exception: If there is an error in processing the image data.
  """

  if attempt > 1:
    return None
  try:
    obj = resolve1(document.getobj(obj_num))
    if obj and isinstance(obj, PDFStream):
      width = obj["Width"]
      height = obj["Height"]
      if width < 5 or height < 5:
        raise Exception("Image too small. Not target image")
      if width > 1000 and height > 1000:
        raise Exception("probably full page image")
      stream = obj.get_data()
      return create_image_from_binary(stream, width, height)
  except Exception as e:
    info_logger.info(f"Issue: {e} with object: {obj_num}")
    return extract_image(document, obj_num + 1, attempt + 1)


def parse_img_obj(file_path: str, obj_nums: list) -> str:
  """
  Extracts and processes images from a PDF file, then performs OCR to convert
  images to text.

  This function takes the path of a PDF file and a list of object numbers
  (obj_nums) which are potential image objects within the PDF. It attempts 
  to extract images from these objects and then uses OCR to convert these 
  images into text. The text from all images is concatenated and returned 
  as a single string.

  Arguments:
    file_path (str): The file path to the PDF from which to extract images.
    obj_nums (list): A list of object numbers in the PDF that potentially
    contain images.

  Returns text (str): the concatenated text extracted from all images.
  """

  with open(file_path, 'rb') as f:
    parser = PDFParser(f)
    document = PDFDocument(parser)
    base64_images = []
    for obj_num in obj_nums:
      image = extract_image(document, obj_num, attempt=0)
      base64_images.append(image) if image else None
  text = run_ocr(base64_images) if base64_images else ""
  return text


def process_element(element: Any) -> [Tuple[str, Any], None]:
  """
  Processes a PDF layout element to identify its type and extract relevant
  data.

  This function categorizes a given PDF layout element into types such as
  image, text, or other. For image elements, it returns the object ID. For
  text elements, it returns the extracted text. The function handles
  container elements by recursively processing their children. If an element
  does not match any specific type, it is categorized as "other".

  Arguments:
    element (Any): A PDF layout element from pdfminer.

  Returns: A tuple containing the element type and its relevant data (object
    ID for images or text content for text elements), or None for other types.
  """

  if hasattr(element, "stream"):
    return ("image", element.stream.objid)
  elif isinstance(element, LTText) and not isinstance(element, LTChar):
    return ("text", element.get_text())
  elif isinstance(element, LTContainer):
    for child in element:
      return process_element(child)
  return ("other", None)


def parse_pdf_page(page: str, metadata: dict) -> str:
  """
  Parses the given pdf page and returns it as a string.
  Arguments:
    page: A pdf page.
  Returns the pdf page as a string.
  """

  page_list = []
  paragraph_list = []

  def parse_paragraph(paragraph_list: list) -> str:
    """
    Check if the first three lines of text contain a identifier for a page
    unrelated to the chapter, or a chapter heading.
    If not a chapter, return an empty list. If a chapter heading, replace with
    three asterisks. Otherwise concantate line to paragraph string and add to
    page list
    """

    text_lines = 0
    paragraph = ""

    for line in paragraph_list:
      if line.strip() or text_lines < 3:
        if is_not_chapter(paragraph, metadata):
          return ""
        elif is_chapter(paragraph):
          return "\n***\n"
      paragraph += line
      text_lines += 1
    return paragraph

  lines = page.split("\n")

  for line in lines:
    if line.strip():
      paragraph_list.append(desmarten_text(line))
    elif paragraph_list:
      paragraph = parse_paragraph(paragraph_list)
      if paragraph:
        page_list.append(paragraph)
        paragraph_list = []
      else:
        return ""

  if paragraph_list:
    paragraph = parse_paragraph(paragraph_list)
    if paragraph:
      page_list.append(paragraph)
    else:
      return ""

  return "\n".join(page_list)


def read_pdf(file_path: str, metadata: dict) -> str:
  """
  Reads the contents of a pdf file and returns it as a string.
  Arguments:
    file_path: Path to the pdf file.
    metadata: Dictionary containing the title and author of the file.
  Returns the contents of the pdf file as a string.
  """

  full_text = ""
  for page in extract_pages(file_path):
    pdf_text = ""
    obj_nums = []
    for element in page:
      obj_tuple = process_element(element)
      if obj_tuple[0] == "image":
        obj_nums.append(obj_tuple[1])
      elif obj_tuple[0] == "text":
        pdf_text += obj_tuple[1] + "\n" if obj_tuple[
            1] != "No text found" else ""
    ocr_text = parse_img_obj(file_path, obj_nums)
    page_text = ocr_text + "\n" + pdf_text if ocr_text is not None else pdf_text
    pdf_page = parse_pdf_page(page_text, metadata)
    if pdf_page:
      if pdf_page.rstrip().endswith(('.', '!', '?', '."', '!"', '?"')):
        full_text += pdf_page
      else:
        full_text += "\n" + pdf_page
  full_text = re.sub(r"\s+", " ", full_text)
  if full_text.startswith("\n***\n"):
    full_text = full_text.lstrip("\n***\n")

  return full_text


def parse_text_file(book_content: str) -> str:
  """
  Parses the book content, replacing chapter headers with asterisks and applying text standardization.
  Arguments:
    book_content: The entire content of the book as a string.
  Returns the processed book content as a string.
  """
  book_lines = book_content.split("\n")
  parsed_lines = [
      "\n***\n" if is_chapter(line) else desmarten_text(line)
      for line in book_lines
  ]
  return "\n".join(parsed_lines)


def convert_file(file_path: str, metadata: dict) -> str:
  """
  Converts a book to a text file with 3 asterisks for chapter breaks
  Arguments:
    book_name: Name of the book.
    folder_name: Name of the folder containing the book.
  """

  book_content = ""
  folder, book_file = os.path.split(file_path)
  book_file = book_file.replace(" ", "_")
  book_file = book_file.replace("-", "_")
  filename_list = book_file.split(".")

  if len(filename_list) > 1:
    base_name = "_".join(filename_list[:-1])
  else:
    base_name = filename_list[0]
  extension = filename_list[-1].lower()

  if extension == "epub":
    book_content = read_epub(file_path, metadata)
  elif extension == "docx":
    book_content = read_docx(file_path, metadata)
  elif extension == "pdf":
    book_content = read_pdf(file_path, metadata)
  elif extension == "txt" or extension == "text":
    book_content = read_text_file(file_path)
    book_content = parse_text_file(book_content)
  else:
    raise ValueError("Invalid file type")

  book_name = f"{base_name}.txt"
  book_path = os.path.join(folder, book_name)
  write_to_file(book_content, book_path)
  return book_path
